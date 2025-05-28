# ingest.py
"""
Document ingestion and processing
"""
import json
import os
from typing import List
from langchain.schema import Document
from vectorstore_utils import create_vectorstore

# Additional imports for new file types
from docx import Document as DocxDocument
import PyPDF2

def load_chunks_from_json(file_path: str) -> List[Document]:
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return []
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, list):
            return [Document(page_content=chunk.get("page_content", ""),
                            metadata=chunk.get("metadata", {}))
                    for chunk in data]
        else:
            print(f"Warning: Expected JSON list in {file_path}")
            return []
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON in {file_path}")
        return []
    except Exception as e:
        print(f"Error loading chunks from {file_path}: {e}")
        return []

def load_chunks_from_txt(file_path: str) -> List[Document]:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception as e:
        print(f"Error reading TXT file {file_path}: {e}")
        return []

def load_chunks_from_docx(file_path: str) -> List[Document]:
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        return []

def load_chunks_from_pdf(file_path: str) -> List[Document]:
    try:
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
        return []

def load_chunks_from_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".json":
        return load_chunks_from_json(file_path)
    elif ext == ".txt":
        return load_chunks_from_txt(file_path)
    elif ext == ".docx":
        return load_chunks_from_docx(file_path)
    elif ext == ".pdf":
        return load_chunks_from_pdf(file_path)
    else:
        print(f"Unsupported file extension: {ext}")
        return []

def process_documents(input_file: str,
                      index_name: str = "faiss_index",
                      model_name: str = "all-MiniLM-L6-v2") -> bool:
    try:
        chunks = load_chunks_from_file(input_file)
        if not chunks:
            print("No document chunks loaded. Check your input file.")
            return False

        print(f"Loaded {len(chunks)} document chunks")

        vectorstore = create_vectorstore(chunks, index_name, model_name)
        print(f"Successfully created vector index with {len(chunks)} documents")
        print(f"Index saved to '{index_name}' folder")

        return True

    except Exception as e:
        print(f"Error processing documents: {e}")
        return False

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Process documents into vector embeddings")
    parser.add_argument("--input", "-i", required=True,
                        help="Input file (JSON, PDF, DOCX, TXT)")
    parser.add_argument("--output", "-o", default="faiss_index",
                        help="Output folder for the FAISS index")
    parser.add_argument("--model", "-m", default="all-MiniLM-L6-v2",
                        help="SentenceTransformer model to use")

    args = parser.parse_args()

    if process_documents(args.input, args.output, args.model):
        print("Processing completed successfully")
    else:
        print("Processing failed")
